from django.shortcuts import render
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
import os, uuid, imghdr
from django.conf import settings
from django.http import JsonResponse, FileResponse, HttpResponse
import io
import filetype
from django.core.files.uploadedfile import UploadedFile
from rembg import remove
import logging
from rembg.bg import remove
import numpy as np
import sys
import gc
import cv2
from django.views.decorators.csrf import csrf_exempt
from scipy.ndimage import median_filter
from skimage import exposure, restoration
from skimage.restoration import denoise_nl_means, estimate_sigma
from skimage.morphology import remove_small_objects, disk
from skimage.filters import median, unsharp_mask
from scipy.ndimage import gaussian_filter
from .utils import enhance_image, encode_image
import base64
import json
from django.core.files.storage import default_storage
import fitz  # PyMuPDF for PDF operations
import docx
import pandas as pd
import csv
import tempfile
import mimetypes
from .services.watermark_detector import WatermarkDetector

logger = logging.getLogger(__name__)

# Initialize background remover

# Create your views here.
def home(request):
    return render(request, 'index.html')

def crop(request):
    if request.method == 'POST':
        try:
            image_file = request.FILES.get('image')
            if not image_file:
                return JsonResponse({'error': 'No image uploaded'}, status=400)

            # Process image with PIL
            try:
                # Read image
                image_data = image_file.read()
                img = Image.open(io.BytesIO(image_data))
                
                # Convert to RGB if needed
                if img.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', img.size, 'white')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background

                # Get crop data
                try:
                    x = max(0, round(float(request.POST.get('x', 0))))
                    y = max(0, round(float(request.POST.get('y', 0))))
                    width = min(round(float(request.POST.get('width', img.size[0]))), img.size[0])
                    height = min(round(float(request.POST.get('height', img.size[1]))), img.size[1])
                except (ValueError, TypeError):
                    return JsonResponse({'error': 'Invalid crop parameters'}, status=400)

                # Perform crop
                if width <= 0 or height <= 0:
                    return JsonResponse({'error': 'Invalid crop dimensions'}, status=400)

                cropped = img.crop((x, y, x + width, y + height))
                
                # Save result
                unique_id = uuid.uuid4().hex[:8]
                filename = f"cropped_{unique_id}.png"
                output_path = os.path.join(settings.MEDIA_ROOT, filename)
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Save with quality settings
                cropped.save(output_path, 'PNG', quality=95, optimize=True)
                
                try:
                    return JsonResponse({
                        'success': True,
                        'image_url': settings.MEDIA_URL + filename,
                        'filename': filename,
                        'dimensions': {'width': width, 'height': height}
                    })
                except Exception as e:
                    # Clean up saved file on error
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    logger.error(f"Response error: {str(e)}")
                    return JsonResponse({'error': 'Failed to process request'}, status=400)

            except Exception as e:
                logger.error(f"Image processing error: {str(e)}")
                return JsonResponse({'error': 'Failed to process image'}, status=400)
            finally:
                if 'img' in locals(): img.close()
                if 'cropped' in locals(): cropped.close()
                gc.collect()

        except Exception as e:
            logger.error(f"Upload error: {str(e)}")
            return JsonResponse({'error': 'Failed to upload image'}, status=400)

    return render(request, 'tools/crop.html')

def background_removal(request):
    if request.method == 'POST':
        try:
            # Get uploaded image
            image_file = request.FILES.get('image')
            if not image_file:
                return JsonResponse({'error': 'No image uploaded'}, status=400)

            # Get adjustment parameters
            brightness = float(request.POST.get('brightness', 100)) / 100
            contrast = float(request.POST.get('contrast', 100)) / 100
            blur = int(request.POST.get('blur', 0))
            bg_type = request.POST.get('bgType', 'transparent')
            bg_color = request.POST.get('bgColor', '#ffffff')

            # Read image
            image_data = image_file.read()
            img = Image.open(io.BytesIO(image_data))

            # Remove background
            output = remove(img)

            # Apply adjustments
            if brightness != 1:
                enhancer = ImageEnhance.Brightness(output)
                output = enhancer.enhance(brightness)

            if contrast != 1:
                enhancer = ImageEnhance.Contrast(output)
                output = enhancer.enhance(contrast)

            # Handle background options
            if bg_type == 'solid':
                # Create solid color background
                bg = Image.new('RGB', output.size, bg_color)
                bg.paste(output, mask=output.split()[3])
                output = bg
            elif bg_type == 'blur':
                # Convert to CV2 format for blur
                cv_img = np.array(img)
                blurred = cv2.GaussianBlur(cv_img, (blur * 2 + 1, blur * 2 + 1), 0)
                bg = Image.fromarray(blurred)
                bg.paste(output, mask=output.split()[3])
                output = bg

            # Save result
            unique_id = uuid.uuid4().hex[:8]
            filename = f"bg_removed_{unique_id}.png"
            output_path = os.path.join(settings.MEDIA_ROOT, filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            output.save(output_path, 'PNG', quality=95, optimize=True)

            return JsonResponse({
                'success': True,
                'image_url': settings.MEDIA_URL + filename
            })

        except Exception as e:
            logger.error(f"Background removal error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
        finally:
            if 'img' in locals(): img.close()
            if 'output' in locals(): output.close()
            gc.collect()

    return render(request, 'background_removal.html')

def download_image(request, filename):
    try:
        file_path = os.path.join(settings.MEDIA_ROOT, filename)
        if os.path.exists(file_path):
            response = FileResponse(open(file_path, 'rb'), content_type='image/png')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        return JsonResponse({'error': 'File not found'}, status=404)
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        return JsonResponse({'error': 'Failed to download image'}, status=400)

def image_compression(request):
    if request.method == 'POST':
        try:
            image_file = request.FILES.get('image')
            if not image_file:
                logger.error("No image file received")
                return JsonResponse({'success': False, 'error': 'No image uploaded'}, status=400)

            try:
                quality = int(request.POST.get('quality', 75))
            except (ValueError, TypeError):
                logger.error("Invalid quality value")
                return JsonResponse({'success': False, 'error': 'Invalid quality value'}, status=400)

            try:
                # Read original image
                image_data = image_file.read()
                original_size = len(image_data)
                img = Image.open(io.BytesIO(image_data))

                # Convert to RGB if needed
                if img.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img.size, 'white')
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background

                # Smart dimension reduction if image is too large
                max_dimension = 2000
                if img.size[0] > max_dimension or img.size[1] > max_dimension:
                    ratio = min(max_dimension / img.size[0], max_dimension / img.size[1])
                    new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)

                # Prepare output with optimization
                output_buffer = io.BytesIO()
                
                # Smart quality adjustment
                target_size = 500 * 1024  # 500KB target
                if quality >= 95:  # Auto mode
                    quality = 95
                    while quality > 5:
                        output_buffer.seek(0)
                        output_buffer.truncate()
                        img.save(output_buffer, format='JPEG', quality=quality, optimize=True)
                        if len(output_buffer.getvalue()) <= target_size:
                            break
                        quality -= 5
                else:
                    img.save(output_buffer, format='JPEG', quality=quality, optimize=True)

                # Save result
                unique_id = uuid.uuid4().hex[:8]
                filename = f"compressed_{unique_id}.jpg"
                output_path = os.path.join(settings.MEDIA_ROOT, filename)
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                with open(output_path, 'wb') as f:
                    f.write(output_buffer.getvalue())

                compressed_size = len(output_buffer.getvalue())
                compression_ratio = ((original_size - compressed_size) / original_size) * 100

                return JsonResponse({
                    'success': True,
                    'image_url': settings.MEDIA_URL + filename,
                    'original_size': original_size,
                    'compressed_size': compressed_size,
                    'compression_ratio': round(compression_ratio, 2),
                    'quality_used': quality,
                    'dimensions': {'width': img.size[0], 'height': img.size[1]}
                }, status=200)

            except Exception as e:
                logger.error(f"Image processing error: {str(e)}")
                return JsonResponse({'success': False, 'error': 'Failed to process image'}, status=400)
            finally:
                if 'img' in locals(): img.close()
                if 'output_buffer' in locals(): output_buffer.close()
                gc.collect()

        except Exception as e:
            logger.error(f"Compression error: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Failed to compress image'}, status=400)

    return render(request, 'tools/image_compression.html')

def passport_photo(request):
    if request.method == 'POST':
        try:
            # Get uploaded image
            image_file = request.FILES.get('image')
            if not image_file:
                return JsonResponse({'error': 'No image uploaded'}, status=400)

            # Get parameters
            bg_color = request.POST.get('bgColor', '#FFFFFF')
            size_option = request.POST.get('size', '2x2')

            # Read image
            image_data = image_file.read()
            img = Image.open(io.BytesIO(image_data))

            # Convert to RGB if needed
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Remove background
            output = remove(img)

            # Get adjustment parameters
            brightness = float(request.POST.get('brightness', 100)) / 100
            contrast = float(request.POST.get('contrast', 100)) / 100
            rotation = int(request.POST.get('rotation', 0))

            # Apply rotation if needed
            if rotation != 0:
                output = output.rotate(rotation, expand=True)

            # Apply adjustments
            if brightness != 1:
                enhancer = ImageEnhance.Brightness(output)
                output = enhancer.enhance(brightness)

            if contrast != 1:
                enhancer = ImageEnhance.Contrast(output)
                output = enhancer.enhance(contrast)

            # Create new image with selected background color
            bg = Image.new('RGB', output.size, bg_color)
            bg.paste(output, mask=output.split()[3])
            output = bg

            # Resize based on selected size
            size_map = {
                '2x2': (600, 600),          # 2x2 inch at 300 DPI
                '35x45': (420, 540),        # 35x45 mm at 300 DPI
                '3.5x4.5': (420, 540)       # 3.5x4.5 cm converted to pixels
            }
            
            target_size = size_map.get(size_option, (600, 600))
            
            # Calculate resize ratio maintaining aspect ratio
            ratio = min(target_size[0]/output.size[0], target_size[1]/output.size[1])
            new_size = tuple(int(dim * ratio) for dim in output.size)
            resized = output.resize(new_size, Image.Resampling.LANCZOS)

            # Create final image with padding to match target size
            final = Image.new('RGB', target_size, bg_color)
            paste_x = (target_size[0] - new_size[0]) // 2
            paste_y = (target_size[1] - new_size[1]) // 2
            final.paste(resized, (paste_x, paste_y))

            # Save result
            unique_id = uuid.uuid4().hex[:8]
            filename = f"passport_{unique_id}.jpg"
            output_path = os.path.join(settings.MEDIA_ROOT, filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            final.save(output_path, 'JPEG', quality=95)

            return JsonResponse({
                'success': True,
                'image_url': settings.MEDIA_URL + filename,
                'filename': filename
            })

        except Exception as e:
            logger.error(f"Passport photo error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
        finally:
            if 'img' in locals(): img.close()
            if 'output' in locals(): output.close()
            if 'final' in locals(): final.close()
            gc.collect()

    return render(request, 'tools/passport.html')

def enhancer(request):
    if request.method == 'POST':
        try:
            # Get uploaded image
            image_file = request.FILES.get('image')
            if not image_file:
                return JsonResponse({'error': 'No image uploaded'}, status=400)

            # Get basic adjustments
            brightness = float(request.POST.get('brightness', 0)) / 100 + 1
            contrast = float(request.POST.get('contrast', 0)) / 100 + 1
            saturation = float(request.POST.get('saturation', 0)) / 100 + 1

            # Get color balance
            red_balance = float(request.POST.get('redBalance', 0)) / 100
            green_balance = float(request.POST.get('greenBalance', 0)) / 100
            blue_balance = float(request.POST.get('blueBalance', 0)) / 100
            color_temp = float(request.POST.get('colorTemp', 0)) / 100
            tint = float(request.POST.get('tint', 0)) / 100

            # Get selected effect
            effect = request.POST.get('effect', '')

            # Process image
            img = Image.open(image_file)
            
            # Convert to RGB if needed
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Apply basic adjustments
            if brightness != 1:
                enhancer = ImageEnhance.Brightness(img)
                img = enhancer.enhance(brightness)
            
            if contrast != 1:
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(contrast)
            
            if saturation != 1:
                enhancer = ImageEnhance.Color(img)
                img = enhancer.enhance(saturation)

            # Apply color balance adjustments
            if any([red_balance, green_balance, blue_balance, color_temp, tint]):
                r, g, b = img.split()
                
                # RGB balance adjustments
                if red_balance:
                    r = ImageEnhance.Brightness(r).enhance(1 + red_balance)
                if green_balance:
                    g = ImageEnhance.Brightness(g).enhance(1 + green_balance)
                if blue_balance:
                    b = ImageEnhance.Brightness(b).enhance(1 + blue_balance)

                # Color temperature
                if color_temp:
                    if color_temp > 0:  # Warmer
                        r = ImageEnhance.Brightness(r).enhance(1 + color_temp)
                        b = ImageEnhance.Brightness(b).enhance(1 - color_temp * 0.8)
                    else:  # Cooler
                        b = ImageEnhance.Brightness(b).enhance(1 - color_temp)
                        r = ImageEnhance.Brightness(r).enhance(1 + color_temp * 0.8)

                # Tint
                if tint:
                    if tint > 0:  # Green tint
                        g = ImageEnhance.Brightness(g).enhance(1 + tint)
                    else:  # Magenta tint
                        r = ImageEnhance.Brightness(r).enhance(1 - tint * 0.8)
                        b = ImageEnhance.Brightness(b).enhance(1 - tint * 0.8)

                img = Image.merge('RGB', (r, g, b))

            # Apply effects
            if effect:
                if effect == 'hdr':
                    img = ImageEnhance.Contrast(img).enhance(1.4)
                    img = ImageEnhance.Color(img).enhance(1.2)
                elif effect == 'cinematic':
                    # Add cinematic wide-screen effect
                    width, height = img.size
                    bar_height = int(height * 0.125)  # 12.5% of height for bars
                    img = ImageOps.crop(img, (0, bar_height, 0, bar_height))
                    img = ImageEnhance.Contrast(img).enhance(1.3)
                    img = ImageEnhance.Color(img).enhance(0.85)
                elif effect == 'vintage':
                    # Add vintage effect with color tinting
                    img = ImageEnhance.Color(img).enhance(0.7)
                    r, g, b = img.split()
                    r = ImageEnhance.Brightness(r).enhance(1.1)
                    b = ImageEnhance.Brightness(b).enhance(0.9)
                    img = Image.merge('RGB', (r, g, b))
                    img = img.filter(ImageFilter.SMOOTH)
                elif effect == 'polaroid':
                    # Create polaroid-style effect
                    img = ImageEnhance.Brightness(img).enhance(1.2)
                    img = ImageEnhance.Contrast(img).enhance(1.1)
                    img = ImageEnhance.Color(img).enhance(1.3)
                    frame = Image.new('RGB', (img.width + 40, img.height + 100), 'white')
                    frame.paste(img, (20, 20))
                    img = frame
                elif effect == 'sunset':
                    # Create warm sunset effect
                    img = ImageEnhance.Color(img).enhance(1.4)
                    r, g, b = img.split()
                    r = ImageEnhance.Brightness(r).enhance(1.4)
                    b = ImageEnhance.Brightness(b).enhance(0.8)
                    img = Image.merge('RGB', (r, g, b))
                elif effect == 'moonlight':
                    # Create cool moonlight effect
                    img = ImageEnhance.Brightness(img).enhance(0.8)
                    img = ImageEnhance.Contrast(img).enhance(1.2)
                    r, g, b = img.split()
                    b = ImageEnhance.Brightness(b).enhance(1.3)
                    img = Image.merge('RGB', (r, g, b))
                    img = img.filter(ImageFilter.GaussianBlur(1))
                elif effect == 'cyberpunk':
                    # Create cyberpunk-style effect
                    img = ImageEnhance.Contrast(img).enhance(1.5)
                    img = ImageEnhance.Color(img).enhance(1.8)
                    r, g, b = img.split()
                    r = ImageEnhance.Brightness(r).enhance(1.4)
                    b = ImageEnhance.Brightness(b).enhance(1.5)
                    img = Image.merge('RGB', (r, g, b))
                elif effect == 'pastels':
                    # Create soft pastel effect
                    img = ImageEnhance.Color(img).enhance(0.7)
                    img = ImageEnhance.Brightness(img).enhance(1.2)
                    img = img.filter(ImageFilter.SMOOTH)
                elif effect == 'infrared':
                    # Simulate infrared photography
                    img = ImageOps.grayscale(img)
                    img = ImageEnhance.Contrast(img).enhance(1.6)
                    img = ImageEnhance.Brightness(img).enhance(1.1)
                    img = ImageOps.colorize(img, 'darkred', 'white')
                elif effect == 'cross_process':
                    # Cross processing effect
                    r, g, b = img.split()
                    r = ImageEnhance.Contrast(r).enhance(1.5)
                    g = ImageEnhance.Brightness(g).enhance(1.1)
                    b = ImageEnhance.Brightness(b).enhance(0.9)
                    img = Image.merge('RGB', (r, g, b))
                elif effect == 'lomo':
                    # Lomo camera effect
                    img = ImageEnhance.Color(img).enhance(1.3)  # Changed from Saturation to Color
                    img = ImageEnhance.Contrast(img).enhance(1.5)
                    r, g, b = img.split()
                    r = ImageEnhance.Brightness(r).enhance(1.2)
                    img = Image.merge('RGB', (r, g, b))
                    # Add vignette
                    width, height = img.size
                    mask = Image.new('L', (width, height))
                    for y in range(height):
                        for x in range(width):
                            distance = ((x - width/2)**2 + (y - height/2)**2)**0.5
                            radius = min(width, height)/1.5
                            value = int(255 * (1 - min(1, distance/radius)))
                            mask.putpixel((x, y), value)
                    img = Image.composite(img, Image.new('RGB', img.size, 'black'), mask)
                elif effect == 'bw':
                    img = ImageOps.grayscale(img).convert('RGB')
                elif effect == 'sepia':
                    img = ImageOps.sepia(img)
                elif effect == 'vivid':
                    img = ImageEnhance.Color(img).enhance(1.5)
                    img = ImageEnhance.Contrast(img).enhance(1.2)
                elif effect == 'soft':
                    img = img.filter(ImageFilter.GaussianBlur(2))
                    img = ImageEnhance.Brightness(img).enhance(1.1)
                elif effect == 'sharp':
                    img = img.filter(ImageFilter.UnsharpMask(radius=2, percent=150))
                elif effect == 'noise':
                    img = img.filter(ImageFilter.MedianFilter(3))
                elif effect == 'film':
                    img = ImageEnhance.Color(img).enhance(0.8)
                    img = ImageEnhance.Contrast(img).enhance(1.2)
                elif effect == 'dramatic':
                    img = ImageEnhance.Contrast(img).enhance(1.5)
                    img = ImageEnhance.Color(img).enhance(0.8)
                elif effect == 'cool':
                    r, g, b = img.split()
                    b = ImageEnhance.Brightness(b).enhance(1.2)
                    img = Image.merge('RGB', (r, g, b))
                elif effect == 'warm':
                    r, g, b = img.split()
                    r = ImageEnhance.Brightness(r).enhance(1.2)
                    img = Image.merge('RGB', (r, g, b))
                elif effect == 'dehaze':
                    img = ImageEnhance.Contrast(img).enhance(1.3)
                    img = ImageEnhance.Brightness(img).enhance(1.1)
                elif effect == 'vignette':
                    # Create vignette effect
                    width, height = img.size
                    mask = Image.new('L', (width, height))
                    for y in range(height):
                        for x in range(width):
                            distance = ((x - width/2)**2 + (y - height/2)**2)**0.5
                            radius = min(width, height)/2
                            mask.putpixel((x, y), int(255 * (1 - min(1, distance/radius))))
                    img = Image.composite(img, Image.new('RGB', img.size, 'black'), mask)

            # Save enhanced image
            unique_id = uuid.uuid4().hex[:8]
            filename = f"enhanced_{unique_id}.jpg"
            output_path = os.path.join(settings.MEDIA_ROOT, filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            img.save(output_path, 'JPEG', quality=95)

            # Add file size to response
            file_size = os.path.getsize(output_path)
            
            return JsonResponse({
                'success': True,
                'image_url': settings.MEDIA_URL + filename,
                'filename': filename,
                'file_size': file_size
            })

        except Exception as e:
            logger.error(f"Image enhancement error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
        finally:
            if 'img' in locals(): img.close()
            gc.collect()

    return render(request, 'tools/enhancer.html')

@csrf_exempt
def restorePhoto(request):
    if request.method == 'POST':
        try:
            image_file = request.FILES.get('image')
            if not image_file:
                return JsonResponse({'error': 'No image uploaded'}, status=400)

            # Create restore directories if they don't exist
            restore_upload_dir = os.path.join(settings.MEDIA_ROOT, 'restore', 'uploads')
            os.makedirs(restore_upload_dir, exist_ok=True)

            # Save original uploaded image
            original_filename = f"original_{uuid.uuid4().hex[:8]}.jpg"
            original_path = os.path.join(restore_upload_dir, original_filename)

            # Read image
            image_data = image_file.read()
            img = Image.open(io.BytesIO(image_data))

            # Convert to RGB if needed
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Convert to numpy array
            img_array = np.array(img)

            # 1. Advanced Noise Reduction with fixed parameters for newer scikit-image versions
            try:
                # First attempt with channel_axis (newer versions)
                sigma_est = np.mean(estimate_sigma(img_array, channel_axis=-1))
                patch_kw = dict(patch_size=5, patch_distance=6, channel_axis=-1)
            except TypeError:
                # Fallback for older versions
                sigma_est = np.mean(estimate_sigma(img_array, multichannel=True))
                patch_kw = dict(patch_size=5, patch_distance=6, multichannel=True)

            denoised = denoise_nl_means(img_array, h=1.15 * sigma_est, fast_mode=True, **patch_kw)
            denoised = (denoised * 255).astype(np.uint8)

            # Alternative denoising if the above fails
            if denoised is None or not isinstance(denoised, np.ndarray):
                denoised = cv2.fastNlMeansDenoisingColored(img_array, None, 10, 10, 7, 21)

            # 2. Enhanced Face Detection and Processing
            face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            gray = cv2.cvtColor((denoised * 255).astype(np.uint8), cv2.COLOR_RGB2GRAY)
            faces = face_cascade.detectMultiScale(gray, 1.3, 5)
            
            face_mask = np.zeros_like(gray)
            for (x, y, w, h) in faces:
                # Expand face region slightly
                exp_x = int(x - w * 0.1)
                exp_y = int(y - h * 0.1)
                exp_w = int(w * 1.2)
                exp_h = int(h * 1.2)
                exp_x = max(0, exp_x)
                exp_y = max(0, exp_y)
                exp_w = min(gray.shape[1] - exp_x, exp_w)
                exp_h = min(gray.shape[0] - exp_y, exp_h)
                face_mask[exp_y:exp_y+exp_h, exp_x:exp_x+exp_w] = 255

            # 3. Advanced Damage Detection
            gradient_x = cv2.Sobel(denoised, cv2.CV_64F, 1, 0, ksize=3)
            gradient_y = cv2.Sobel(denoised, cv2.CV_64F, 0, 1, ksize=3)
            gradient_magnitude = np.sqrt(gradient_x**2 + gradient_y**2)
            damage_mask = (gradient_magnitude > np.percentile(gradient_magnitude, 90)).astype(np.uint8) * 255
            
            # Remove small noise from damage mask
            damage_mask = remove_small_objects(damage_mask.astype(bool), min_size=50).astype(np.uint8) * 255

            # 4. Enhanced Inpainting
            radius = 5  # Increased radius for better restoration
            damaged = cv2.inpaint(denoised.astype(np.uint8), damage_mask.astype(np.uint8), 
                                radius, cv2.INPAINT_TELEA)

            # 5. Advanced Color Correction
            lab = cv2.cvtColor(damaged, cv2.COLOR_RGB2LAB)
            l, a, b = cv2.split(lab)
            
            # Enhanced CLAHE with optimal parameters
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            l = clahe.apply(l)
            
            # Color balance correction
            a = cv2.normalize(a, None, 0, 255, cv2.NORM_MINMAX)
            b = cv2.normalize(b, None, 0, 255, cv2.NORM_MINMAX)
            
            enhanced_lab = cv2.merge([l, a, b])
            color_corrected = cv2.cvtColor(enhanced_lab, cv2.COLOR_LAB2RGB)

            # 6. Detail Enhancement with Edge Preservation
            bilateral = cv2.bilateralFilter(color_corrected, 9, 75, 75)
            detail_mask = unsharp_mask(bilateral, radius=2, amount=2)
            
            # 7. Smart Sharpening
            face_regions = face_mask > 0
            result = np.copy(detail_mask)
            
            # Apply stronger sharpening to non-face regions
            non_face_regions = ~face_regions
            result[non_face_regions] = unsharp_mask(detail_mask, radius=1, amount=2)[non_face_regions]
            
            # Gentler sharpening for face regions
            result[face_regions] = unsharp_mask(detail_mask, radius=1, amount=1.5)[face_regions]

            # 8. Final Color Enhancement
            result = exposure.adjust_gamma(result, 0.9)
            
            # Enhance saturation for faded photos
            hsv = cv2.cvtColor((result * 255).astype(np.uint8), cv2.COLOR_RGB2HSV)
            h, s, v = cv2.split(hsv)
            s = cv2.multiply(s, 1.2)  # Increase saturation
            merged_hsv = cv2.merge([h, s, v])
            result = cv2.cvtColor(merged_hsv, cv2.COLOR_HSV2RGB)

            # Convert back to PIL Image with enhanced processing
            restored_image = Image.fromarray(result)
            
            # Final adjustments
            enhancer = ImageEnhance.Contrast(restored_image)
            restored_image = enhancer.enhance(1.2)
            enhancer = ImageEnhance.Color(restored_image)
            restored_image = enhancer.enhance(1.1)

            # Save restored image in restore directory
            restored_filename = f"restored_{uuid.uuid4().hex[:8]}.jpg"
            output_path = os.path.join(restore_upload_dir, restored_filename)
            
            restored_image.save(output_path, 'JPEG', quality=95)

            return JsonResponse({
                'success': True,
                'image_url': f"{settings.MEDIA_URL}restore/uploads/{restored_filename}",
                'original_url': f"{settings.MEDIA_URL}restore/uploads/{original_filename}",
                'message': 'Photo restored successfully'
            })

        except Exception as e:
            logger.error(f"Photo restoration error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
        finally:
            if 'img' in locals(): img.close()
            if 'restored_image' in locals(): restored_image.close()
            gc.collect()

    return render(request, 'tools/restore-photo.html')

def aiFilters(request):
    return render(request, 'tools/Ai-filters.html')

def apply_cartoon_effect(img):
    """Apply cartoon effect using edge detection and color quantization"""
    # Convert to grayscale for edge detection
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    # Apply bilateral filter for smoothing while preserving edges
    smooth = cv2.bilateralFilter(img, 9, 75, 75)
    # Detect edges
    edges = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, 
                                cv2.THRESH_BINARY, 9, 2)
    # Combine smooth colors with edges
    cartoon = cv2.bitwise_and(smooth, smooth, mask=edges)
    return cartoon

def apply_oil_painting_effect(img):
    """Create oil painting effect using stylization"""
    # Apply bilateral filter for initial smoothing
    smooth = cv2.bilateralFilter(img, 9, 75, 75)
    # Apply stylization filter
    oil_painting = cv2.stylization(smooth, sigma_s=60, sigma_r=0.6)
    return oil_painting

def apply_watercolor_effect(img):
    """Create watercolor effect using color quantization and blur"""
    # Quantize colors
    k = 8  # number of colors
    data = img.reshape((-1, 3))
    data = np.float32(data)
    criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 20, 1.0)
    _, labels, centers = cv2.kmeans(data, k, None, criteria, 10, cv2.KMEANS_RANDOM_CENTERS)
    centers = np.uint8(centers)
    quantized = centers[labels.flatten()]
    quantized = quantized.reshape(img.shape)
    
    # Apply bilateral filter for watercolor effect
    watercolor = cv2.bilateralFilter(quantized, 9, 75, 75)
    return watercolor

def apply_sketch_effect(img):
    """Create pencil sketch effect"""
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    # Invert grayscale image
    inverted = 255 - gray
    # Apply Gaussian blur
    blurred = cv2.GaussianBlur(inverted, (21, 21), 0)
    # Invert blurred image
    inverted_blurred = 255 - blurred
    # Create pencil sketch
    sketch = cv2.divide(gray, inverted_blurred, scale=256.0)
    # Convert back to RGB
    sketch_rgb = cv2.cvtColor(sketch, cv2.COLOR_GRAY2RGB)
    return sketch_rgb

def apply_vintage_effect(img):
    """Apply vintage/retro effect"""
    # Split channels
    b, g, r = cv2.split(img)
    # Adjust individual channels for vintage look
    r = cv2.normalize(r, None, 100, 255, cv2.NORM_MINMAX)
    b = cv2.normalize(b, None, 0, 255, cv2.NORM_MINMAX)
    # Add sepia tone
    vintage = cv2.merge([
        b * 0.8,
        g * 0.9,
        r * 1.2
    ])
    # Add vignette effect
    rows, cols = img.shape[:2]
    kernel_x = cv2.getGaussianKernel(cols, cols/2)
    kernel_y = cv2.getGaussianKernel(rows, rows/2)
    kernel = kernel_y * kernel_x.T
    mask = 255 * kernel / np.linalg.norm(kernel)
    vintage = vintage * mask[:, :, np.newaxis]
    return vintage

def apply_cyberpunk_effect(img):
    """Create cyberpunk-style effect with neon colors"""
    # Increase contrast
    lab = cv2.cvtColor(img, cv2.COLOR_RGB2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    l = clahe.apply(l)
    lab = cv2.merge([l, a, b])
    img = cv2.cvtColor(lab, cv2.COLOR_LAB2RGB)
    
    # Adjust colors for cyberpunk look
    hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
    h, s, v = cv2.split(hsv)
    s = cv2.normalize(s, None, 100, 255, cv2.NORM_MINMAX)
    v = cv2.normalize(v, None, 100, 255, cv2.NORM_MINMAX)
    hsv = cv2.merge([h, s, v])
    cyberpunk = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
    
    # Add glow effect
    blur = cv2.GaussianBlur(cyberpunk, (0, 0), 10)
    cyberpunk = cv2.addWeighted(cyberpunk, 1.5, blur, -0.5, 0)
    return cyberpunk

def apply_pop_art_effect(img):
    """Create pop art style effect"""
    # Quantize to fewer colors
    Z = img.reshape((-1, 3))
    Z = np.float32(Z)
    criteria = (cv2.TERM_CRITERIA_EPS + cv2.TERM_CRITERIA_MAX_ITER, 10, 1.0)
    K = 8
    _, labels, centers = cv2.kmeans(Z, K, None, criteria, 10, cv2.KMEANS_RANDOM_CENTERS)
    centers = np.uint8(centers)
    res = centers[labels.flatten()]
    pop_art = res.reshape(img.shape)
    
    # Enhance colors
    hsv = cv2.cvtColor(pop_art, cv2.COLOR_RGB2HSV)
    h, s, v = cv2.split(hsv)
    s = cv2.normalize(s, None, 150, 255, cv2.NORM_MINMAX)
    hsv = cv2.merge([h, s, v])
    pop_art = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
    return pop_art

def apply_hdr_effect(img):
    """Create HDR-like effect"""
    # Convert to LAB color space
    lab = cv2.cvtColor(img, cv2.COLOR_RGB2LAB)
    l, a, b = cv2.split(lab)
    
    # Apply CLAHE to L channel
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    l = clahe.apply(l)
    
    # Merge channels
    lab = cv2.merge([l, a, b])
    
    # Convert back to RGB
    hdr = cv2.cvtColor(lab, cv2.COLOR_LAB2RGB)
    
    # Enhance contrast and saturation
    hsv = cv2.cvtColor(hdr, cv2.COLOR_RGB2HSV)
    h, s, v = cv2.split(hsv)
    s = cv2.normalize(s, None, 100, 255, cv2.NORM_MINMAX)
    hsv = cv2.merge([h, s, v])
    hdr = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
    
    return hdr

def apply_basic_filter(img, filter_type):
    """Apply basic color and tone adjustments"""
    if filter_type in ['cool-blue', 'moonlight']:
        # Cool blue tint
        b, g, r = cv2.split(img)
        b = cv2.normalize(b, None, 100, 255, cv2.NORM_MINMAX)
        return cv2.merge([b * 1.2, g * 0.9, r * 0.8])
    
    elif filter_type in ['warm-gold', 'sunset']:
        # Warm golden tint
        b, g, r = cv2.split(img)
        r = cv2.normalize(r, None, 100, 255, cv2.NORM_MINMAX)
        return cv2.merge([b * 0.8, g * 0.9, r * 1.2])
    
    elif filter_type == 'dramatic':
        # High contrast dramatic look
        lab = cv2.cvtColor(img, cv2.COLOR_RGB2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        l = clahe.apply(l)
        lab = cv2.merge([l, a, b])
        return cv2.cvtColor(lab, cv2.COLOR_LAB2RGB)
    
    else:
        # Default enhancement
        return exposure.adjust_gamma(img, 0.9)

def apply_artistic_filter(img, filter_type):
    """Apply artistic filters with advanced processing"""
    if filter_type == 'impressionist':
        # Create impressionist painting effect
        blurred = cv2.GaussianBlur(img, (21, 21), 0)
        enhanced = cv2.addWeighted(img, 2.5, blurred, -1.5, 0)
        return enhanced

    elif filter_type == 'pointillism':
        # Create pointillism effect
        height, width = img.shape[:2]
        points = np.zeros_like(img)
        num_points = (height * width) // 100
        
        for _ in range(num_points):
            x = np.random.randint(0, width)
            y = np.random.randint(0, height)
            color = img[y, x]
            cv2.circle(points, (x, y), 2, color.tolist(), -1)
        
        return points

    elif filter_type == 'cubism':
        # Create cubism-like effect
        segments = 50
        points = []
        
        for _ in range(segments):
            x = np.random.randint(0, img.shape[1])
            y = np.random.randint(0, img.shape[0])
            points.append((x, y))
        
        subdiv = cv2.Subdiv2D((0, 0, img.shape[1], img.shape[0]))
        for p in points:
            subdiv.insert(p)
            
        triangles = subdiv.getTriangleList()
        triangles = np.array(triangles, dtype=np.int32)
        
        result = np.zeros_like(img)
        for t in triangles:
            pt1 = (t[0], t[1])
            pt2 = (t[2], t[3])
            pt3 = (t[4], t[5])
            
            mask = np.zeros((img.shape[0], img.shape[1]), dtype=np.uint8)
            cv2.fillConvexPoly(mask, np.array([pt1, pt2, pt3]), 1)
            
            color = np.mean(img[mask > 0], axis=0)
            cv2.fillConvexPoly(result, np.array([pt1, pt2, pt3]), color.tolist())
            
        return result

def apply_color_effect(img, filter_type):
    """Apply color effects and transformations"""
    if filter_type == 'duotone':
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        color1 = np.array([66, 135, 245])  # Blue
        color2 = np.array([245, 66, 66])   # Red
        
        normalized = gray / 255.0
        result = np.zeros_like(img, dtype=np.float32)
        
        for i in range(3):
            result[:,:,i] = normalized * color1[i] + (1 - normalized) * color2[i]
            
        return np.clip(result, 0, 255).astype(np.uint8)

    elif filter_type == 'neon':
        # Create neon glow effect
        blurred = cv2.GaussianBlur(img, (0, 0), 3)
        edges = cv2.Canny(blurred, 100, 200)
        edges = cv2.cvtColor(edges, cv2.COLOR_GRAY2RGB)
        
        # Enhance colors
        hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
        h, s, v = cv2.split(hsv)
        s = cv2.add(s, 50)
        v = cv2.add(v, 50)
        hsv = cv2.merge([h, s, v])
        colored = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
        
        # Add glow
        glow = cv2.addWeighted(colored, 1.2, edges, 0.8, 0)
        return glow

def apply_texture_effect(img, filter_type):
    """Apply texture effects to images"""
    if filter_type in ['paper', 'canvas', 'metal', 'wood', 'leather', 'fabric']:
        # Load texture overlay
        texture_path = f'media/textures/{filter_type}.jpg'  # You'll need to provide these textures
        try:
            texture = cv2.imread(texture_path)
            texture = cv2.resize(texture, (img.shape[1], img.shape[0]))
            
            # Blend texture with image
            result = cv2.addWeighted(img, 0.8, texture, 0.2, 0)
            return result
        except:
            # Fallback if texture file not found
            return apply_basic_texture(img, filter_type)
    return img

def apply_basic_texture(img, filter_type):
    """Apply procedural textures when texture files are not available"""
    height, width = img.shape[:2]
    texture = np.zeros((height, width, 3), dtype=np.uint8)
    
    if filter_type == 'paper':
        # Create paper-like texture
        noise = np.random.normal(0, 25, (height, width))
        texture = np.clip(img + noise[:, :, np.newaxis], 0, 255).astype(np.uint8)
        
    elif filter_type == 'canvas':
        # Create canvas-like texture
        for i in range(0, height, 4):
            for j in range(0, width, 4):
                noise = np.random.normal(0, 20)
                texture[i:i+4, j:j+4] = noise
        texture = cv2.GaussianBlur(texture, (3, 3), 0)
        result = cv2.addWeighted(img, 0.9, texture, 0.1, 0)
        return result
        
    return cv2.addWeighted(img, 0.85, texture, 0.15, 0)

def apply_weather_effect(img, filter_type):
    """Apply weather and environmental effects"""
    if filter_type == 'rain':
        rain_layer = np.zeros_like(img)
        num_drops = 1000
        
        for _ in range(num_drops):
            x1 = np.random.randint(0, img.shape[1])
            y1 = np.random.randint(0, img.shape[0])
            length = np.random.randint(5, 15)
            angle = -30  # Rain angle in degrees
            x2 = int(x1 + length * np.cos(np.radians(angle)))
            y2 = int(y1 + length * np.sin(np.radians(angle)))
            cv2.line(rain_layer, (x1, y1), (x2, y2), (200, 200, 200), 1)
            
        # Add blur to rain
        rain_layer = cv2.GaussianBlur(rain_layer, (3, 3), 0)
        result = cv2.addWeighted(img, 0.8, rain_layer, 0.2, 0)
        return result

    elif filter_type == 'snow':
        snow_layer = np.zeros_like(img)
        num_flakes = 1000
        
        for _ in range(num_flakes):
            x = np.random.randint(0, img.shape[1])
            y = np.random.randint(0, img.shape[0])
            size = np.random.randint(1, 4)
            cv2.circle(snow_layer, (x, y), size, (255, 255, 255), -1)
            
        # Add blur to snow
        snow_layer = cv2.GaussianBlur(snow_layer, (3, 3), 0)
        result = cv2.addWeighted(img, 0.8, snow_layer, 0.2, 0)
        return result

    elif filter_type in ['fog', 'mist']:
        fog_layer = np.ones_like(img) * 255
        intensity = 0.4 if filter_type == 'fog' else 0.3
        return cv2.addWeighted(img, 1 - intensity, fog_layer, intensity, 0)

def apply_time_effect(img, filter_type):
    """Apply time-of-day effects"""
    if filter_type == 'day':
        # Increase brightness and enhance colors
        enhanced = cv2.convertScaleAbs(img, alpha=1.2, beta=10)
        return enhanced

    elif filter_type == 'night':
        # Reduce brightness and add blue tint
        darkened = cv2.convertScaleAbs(img, alpha=0.6, beta=-20)
        b, g, r = cv2.split(darkened)
        b = cv2.add(b, 20)
        return cv2.merge([b, g, r])

    elif filter_type == 'golden-hour':
        # Add warm golden tones
        hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
        h, s, v = cv2.split(hsv)
        s = cv2.add(s, 30)
        v = cv2.add(v, 10)
        hsv = cv2.merge([h, s, v])
        enhanced = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
        
        # Add warm color overlay
        overlay = np.ones_like(img) * [20, 40, 80]
        return cv2.addWeighted(enhanced, 0.8, overlay, 0.2, 0)

def apply_modern_filter(img, filter_type):
    """Apply modern and trending filters"""
    if filter_type == 'instagram':
        # Create Instagram-like filter
        contrast = cv2.convertScaleAbs(img, alpha=1.2, beta=0)
        hsv = cv2.cvtColor(contrast, cv2.COLOR_RGB2HSV)
        h, s, v = cv2.split(hsv)
        s = cv2.multiply(s, 1.2)
        hsv = cv2.merge([h, s, v])
        return cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)

    elif filter_type == 'minimalist':
        # Create minimalist effect
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        edges = cv2.Canny(gray, 100, 200)
        return cv2.cvtColor(edges, cv2.COLOR_GRAY2RGB)

    elif filter_type == 'retrowave':
        # Create retrowave aesthetic
        hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
        h, s, v = cv2.split(hsv)
        h = (h + 120) % 180
        s = cv2.add(s, 50)
        hsv = cv2.merge([h, s, v])
        enhanced = cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
        
        # Add gradient overlay
        height, width = img.shape[:2]
        gradient = np.zeros((height, width, 3), dtype=np.uint8)
        for i in range(height):
            gradient[i, :] = [
                int(255 * (1 - i/height)),  # Blue
                int(128 * (i/height)),      # Green
                int(255 * (i/height))       # Red
            ]
        return cv2.addWeighted(enhanced, 0.7, gradient, 0.3, 0)

@csrf_exempt
def apply_filter(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            image_data = data.get('image')
            filter_type = data.get('filter')
            
            if not image_data or not filter_type:
                return JsonResponse({'error': 'Missing image or filter type'}, status=400)
            
            # Clean base64 data
            if ',' in image_data:
                image_data = image_data.split(',')[1]
                
            try:
                image_bytes = base64.b64decode(image_data)
                img = Image.open(io.BytesIO(image_bytes))
                
                # Convert to RGB if needed
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Convert to numpy array for processing
                img_array = np.array(img)
                
                # Filter processing based on type
                if filter_type in ['cartoon', 'oil-painting', 'watercolor', 'sketch']:
                    processed = apply_cartoon_effect(img_array) if filter_type == 'cartoon' else \
                               apply_oil_painting_effect(img_array) if filter_type == 'oil-painting' else \
                               apply_watercolor_effect(img_array) if filter_type == 'watercolor' else \
                               apply_sketch_effect(img_array)
                elif filter_type in ['impressionist', 'pointillism', 'cubism']:
                    processed = apply_artistic_filter(img_array, filter_type)
                elif filter_type in ['neon', 'duotone']:
                    processed = apply_color_effect(img_array, filter_type)
                elif filter_type in ['paper', 'canvas', 'metal', 'wood', 'leather', 'fabric']:
                    processed = apply_texture_effect(img_array, filter_type)
                elif filter_type in ['rain', 'snow', 'fog', 'mist']:
                    processed = apply_weather_effect(img_array, filter_type)
                elif filter_type in ['day', 'night', 'golden-hour']:
                    processed = apply_time_effect(img_array, filter_type)
                elif filter_type in ['instagram', 'minimalist', 'retrowave']:
                    processed = apply_modern_filter(img_array, filter_type)
                else:
                    processed = apply_basic_filter(img_array, filter_type)
                
                # Convert back to PIL Image
                processed_img = Image.fromarray(processed.astype('uint8'))
                
                # Save to bytes
                output_buffer = io.BytesIO()
                processed_img.save(output_buffer, format='JPEG', quality=95)
                processed_base64 = base64.b64encode(output_buffer.getvalue()).decode('utf-8')
                
                return JsonResponse({
                    'success': True,
                    'image': f'data:image/jpeg;base64,{processed_base64}'
                })
                
            except Exception as e:
                logger.error(f"Processing error: {str(e)}")
                return JsonResponse({'error': 'Image processing failed'}, status=500)
                
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return JsonResponse({'error': 'An unexpected error occurred'}, status=500)
            
    return JsonResponse({'error': 'Invalid request method'}, status=405)

def formate_convert(request):
    return render(request, "tools/formate-convert.html")

@csrf_exempt
def convert_file(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=400)
    
    try:
        file = request.FILES.get('file')
        target_format = request.POST.get('format')
        
        if not file or not target_format:
            return JsonResponse({'error': 'File and target format are required'}, status=400)
        
        # Validate file size (max 10MB)
        if file.size > 10 * 1024 * 1024:
            return JsonResponse({'error': 'File size exceeds 10MB limit'}, status=400)
        
        # Create a unique filename
        original_name = file.name
        file_extension = os.path.splitext(original_name)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        
        # Create temp directory if it doesn't exist
        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Save the uploaded file temporarily
        temp_path = os.path.join(temp_dir, unique_filename)
        with open(temp_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        
        # Determine file type
        mime_type, _ = mimetypes.guess_type(original_name)
        is_image = mime_type and mime_type.startswith('image/')
        is_document = mime_type and mime_type in [
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'text/plain',
            'text/csv'
        ]
        
        if not (is_image or is_document):
            os.remove(temp_path)
            return JsonResponse({'error': 'Unsupported file type'}, status=400)
        
        # Create output filename and directory
        output_filename = f"{os.path.splitext(original_name)[0]}.{target_format}"
        output_dir = os.path.join(settings.MEDIA_ROOT, 'converted')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        
        # Convert based on file type
        if is_image:
            success = convert_image(temp_path, output_path, target_format)
        else:
            success = convert_document(temp_path, output_path, target_format)
        
        # Clean up temporary file
        try:
            os.remove(temp_path)
        except:
            pass
        
        if not success:
            # Clean up output file if conversion failed
            try:
                if os.path.exists(output_path):
                    os.remove(output_path)
            except:
                pass
            return JsonResponse({'error': 'Conversion failed'}, status=500)
        
        # Return the converted file
        try:
            with open(output_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type=mimetypes.guess_type(output_path)[0])
                response['Content-Disposition'] = f'attachment; filename="{output_filename}"'
                
                # Clean up the converted file after sending
                try:
                    os.remove(output_path)
                except:
                    pass
                    
                return response
        except Exception as e:
            # Clean up output file if sending failed
            try:
                if os.path.exists(output_path):
                    os.remove(output_path)
            except:
                pass
            raise e
            
    except Exception as e:
        logger.error(f"File conversion error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

def convert_image(input_path, output_path, target_format):
    try:
        with Image.open(input_path) as img:
            # Convert RGBA to RGB if necessary
            if img.mode in ('RGBA', 'LA'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            
            # Validate target format
            valid_formats = ['PNG', 'JPEG', 'JPG', 'WEBP', 'GIF', 'BMP', 'TIFF', 'SVG']
            if target_format.upper() not in valid_formats:
                raise ValueError(f"Unsupported image format: {target_format}")
            
            # Save in target format
            img.save(output_path, format=target_format.upper())
            return True
    except Exception as e:
        logger.error(f"Image conversion error: {str(e)}")
        return False

def convert_document(input_path, output_path, target_format):
    try:
        input_ext = os.path.splitext(input_path)[1].lower()
        
        # Validate input and output format combination
        valid_conversions = {
            '.pdf': ['docx', 'txt'],
            '.docx': ['pdf', 'txt'],
            '.doc': ['pdf', 'txt'],
            '.xlsx': ['csv'],
            '.xls': ['csv'],
            '.csv': ['xlsx'],
            '.txt': ['pdf', 'docx']
        }
        
        if input_ext not in valid_conversions or target_format not in valid_conversions[input_ext]:
            raise ValueError(f"Unsupported conversion: {input_ext} to {target_format}")
        
        if input_ext == '.pdf':
            if target_format == 'docx':
                return convert_pdf_to_docx(input_path, output_path)
            elif target_format == 'txt':
                return convert_pdf_to_txt(input_path, output_path)
        
        elif input_ext in ('.docx', '.doc'):
            if target_format == 'pdf':
                return convert_docx_to_pdf(input_path, output_path)
            elif target_format == 'txt':
                return convert_docx_to_txt(input_path, output_path)
        
        elif input_ext in ('.xlsx', '.xls'):
            if target_format == 'csv':
                return convert_excel_to_csv(input_path, output_path)
        
        elif input_ext == '.csv':
            if target_format == 'xlsx':
                return convert_csv_to_excel(input_path, output_path)
        
        elif input_ext == '.txt':
            if target_format in ['pdf', 'docx']:
                return convert_txt_to_docx(input_path, output_path)
        
        return False
    except Exception as e:
        logger.error(f"Document conversion error: {str(e)}")
        return False

def convert_txt_to_docx(input_path, output_path):
    try:
        doc = docx.Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
            doc.add_paragraph(text)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"TXT to DOCX conversion error: {str(e)}")
        return False

def convert_pdf_to_docx(input_path, output_path):
    try:
        doc = fitz.open(input_path)
        docx_doc = docx.Document()
        
        for page in doc:
            text = page.get_text()
            docx_doc.add_paragraph(text)
        
        docx_doc.save(output_path)
        doc.close()
        return True
    except Exception as e:
        print(f"PDF to DOCX conversion error: {str(e)}")
        return False

def convert_pdf_to_txt(input_path, output_path):
    try:
        doc = fitz.open(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for page in doc:
                f.write(page.get_text())
        doc.close()
        return True
    except Exception as e:
        print(f"PDF to TXT conversion error: {str(e)}")
        return False

def convert_docx_to_pdf(input_path, output_path):
    try:
        doc = docx.Document(input_path)
        # Note: This requires additional setup for docx to PDF conversion
        # You might want to use a different library or service for this
        return False
    except Exception as e:
        print(f"DOCX to PDF conversion error: {str(e)}")
        return False

def convert_docx_to_txt(input_path, output_path):
    try:
        doc = docx.Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for paragraph in doc.paragraphs:
                f.write(paragraph.text + '\n')
        return True
    except Exception as e:
        print(f"DOCX to TXT conversion error: {str(e)}")
        return False

def convert_excel_to_csv(input_path, output_path):
    try:
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False)
        return True
    except Exception as e:
        print(f"Excel to CSV conversion error: {str(e)}")
        return False

def convert_csv_to_excel(input_path, output_path):
    try:
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False)
        return True
    except Exception as e:
        print(f"CSV to Excel conversion error: {str(e)}")
        return False

def about(request):
    return render(request, 'about.html')

import json
import uuid
import os
from PIL import Image
from django.conf import settings
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render
from .services.image_processor import ImageProcessor
import logging

logger = logging.getLogger(__name__)

@csrf_exempt
def remove_watermark(request):
    if request.method == 'POST':
        try:
            if request.POST.get('action') == 'detect':
                return detect_watermark(request)
            elif request.POST.get('action') == 'download':
                return handle_download(request)
            return process_image(request)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return render(request, 'tools/remove-watermark.html')

def detect_watermark(request):
    image_file = request.FILES.get('image')
    if not image_file:
        return JsonResponse({'error': 'No image uploaded'}, status=400)

    try:
        img = Image.open(image_file)
        detector = WatermarkDetector()
        watermark = detector.detect(img)
        
        if watermark:
            return JsonResponse({
                'success': True,
                'watermark': watermark
            })
        return JsonResponse({
            'success': False,
            'error': 'No watermark detected'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)

def handle_download(request):
    filename = request.POST.get('filename')
    if not filename:
        return JsonResponse({'error': 'No filename provided'}, status=400)
    
    file_path = os.path.join(settings.MEDIA_ROOT, 'processed', filename)
    if not os.path.exists(file_path):
        return JsonResponse({'error': 'File not found'}, status=404)
        
    with open(file_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='image/png')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

def process_image(request):
    # Validate request
    image_file = request.FILES.get('image')
    if not image_file:
        return JsonResponse({'error': 'No image uploaded'}, status=400)

    # Load and validate image
    try:
        img = Image.open(image_file)
    except Exception as e:
        logger.error(f"Error loading image: {str(e)}")
        return JsonResponse({'error': 'Invalid image file'}, status=400)

    # Get and validate mask coordinates
    try:
        mask_data = json.loads(request.POST.get('mask', '{}'))
        scale_factor = float(request.POST.get('scale_factor', 1.0))
        
        scaled_mask = {
            'x': int(mask_data['x'] * scale_factor),
            'y': int(mask_data['y'] * scale_factor),
            'width': int(mask_data['width'] * scale_factor),
            'height': int(mask_data['height'] * scale_factor)
        }
        
        # Validate coordinates
        if any(v < 0 for v in scaled_mask.values()):
            raise ValueError("Invalid mask coordinates")
            
        if scaled_mask['x'] + scaled_mask['width'] > img.width or \
           scaled_mask['y'] + scaled_mask['height'] > img.height:
            raise ValueError("Mask coordinates outside image bounds")

    except Exception as e:
        logger.error(f"Error processing mask coordinates: {str(e)}")
        return JsonResponse({'error': 'Invalid mask coordinates'}, status=400)

    # Process image
    try:
        processor = ImageProcessor()
        result_img = processor.remove_watermark(img, scaled_mask)
    except Exception as e:
        logger.error(f"Error removing watermark: {str(e)}")
        return JsonResponse({'error': 'Failed to remove watermark'}, status=400)

    # Save result
    try:
        output_dir = os.path.join(settings.MEDIA_ROOT, 'processed')
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"watermark_removed_{uuid.uuid4().hex[:8]}.png"
        output_path = os.path.join(output_dir, filename)
        
        result_img.save(output_path, 'PNG', quality=95)

        return JsonResponse({
            'success': True,
            'image_url': settings.MEDIA_URL + 'processed/' + filename,
            'filename': filename
        })
    except Exception as e:
        logger.error(f"Error saving result: {str(e)}")
        return JsonResponse({'error': 'Failed to save result'}, status=400)

def error_404(request, exception):
    return render(request, 'errors/notfound-404.html', status=404)

def error_500(request):
    return render(request, 'errors/servererror-500.html', status=500)

def error_403(request, exception):
    return render(request, 'errors/forbidden-403.html', status=403)
    return render(request, 'errors/forbidden-403.html', status=403)